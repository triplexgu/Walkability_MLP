globals().clear()
import geopandas as gpd
import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns

def transfer_word(df,doc_name):
    #pip install python-docx
    from docx import Document
    from docx.enum.table import WD_ALIGN_VERTICAL
    # Create a new Word document
    doc = Document()
    doc.add_paragraph('DataFrame Content:')
    table = doc.add_table(df.shape[0] + 1, df.shape[1])
    table.style = 'Table Grid'
    # Add headers to the first row
    for j in range(df.shape[1]):
        table.cell(0, j).text = df.columns[j]

    # Add the data from the DataFrame
    for i in range(df.shape[0]):
        for j in range(df.shape[1]):
            table.cell(i + 1, j).text = str(df.values[i, j])
    # Save the Word document
    doc.save(r"D:\WUR\master\MLP\master thesis\data\数据汇总\comparison_analysis\%s.docx"% doc_name)


import matplotlib.pyplot as plt


def create_box(df1, df2, df3, var):
    # Calculate statistics for each site
    site1_stats = ['1092', df1[var].quantile(0.25), df1[var].quantile(0.75), df1[var].mean()]
    site2_stats = ['1094', df2[var].quantile(0.25), df2[var].quantile(0.75), df2[var].mean()]
    site3_stats = ['two sites', df3[var].quantile(0.25), df3[var].quantile(0.75), df3[var].mean()]

    # Combine the statistics into a single DataFrame
    combined_stats = pd.DataFrame([site1_stats, site2_stats, site3_stats], columns=['Site', '25%', '75%', 'mean'])

    # Set Seaborn style and font scale
    sns.set(style="whitegrid", font_scale=1.5)

    # Create a box plot
    plt.figure(figsize=(32, 30))  # Increase the figure size for more space

    # Clean and professional box plot
    ax = sns.boxplot(x='Site', y=var,
                     data=pd.concat([df1.assign(Site='1092'), df2.assign(Site='1094'), df3.assign(Site='two sites')]),
                     color='white',
                     linewidth=3,  # Box border width
                     fliersize=5,  # Size of outlier points
                     notch=False,  # Set to True for notched box plot
                     showfliers=True,  # Show outlier points
                     medianprops=dict(color='blue', linewidth=3),  # Median line properties
                     boxprops=dict(linewidth=3, edgecolor='black'),  # Box outline properties
                     whiskerprops=dict(linewidth=3, color='black'),  # Whisker properties
                     capprops=dict(color='black', linewidth=3)  # Caps on the whiskers properties
                     )

    # Add lines and labels for the statistics
    for i, site in enumerate(combined_stats['Site']):
        plt.text(i, ax.get_ylim()[0] - 0.05, f'Mean: {combined_stats.iloc[i, 3]:.2f}',
                 horizontalalignment='center', verticalalignment='bottom', color='blue', fontsize=25)
        for j, stat in enumerate(['25%', '75%']):
            plt.text(i, combined_stats.iloc[i, j + 1], f'{stat}: {combined_stats.iloc[i, j + 1]:.2f}',
                     horizontalalignment='left', verticalalignment='bottom', color='blue', fontsize=25)

    # Stylish layout grids
    sns.set(style="whitegrid", rc={"grid.linestyle": "--"})
    # Adjust x-axis tick label properties
    plt.yticks(fontsize=23)  # Increase font size
    plt.xticks(fontsize=25)#Increase font size
    plt.tick_params(axis='x', which='major', pad=20)  # Add more space between the notes and the axis

    plt.xlabel('Site', fontsize=25)
    plt.ylabel('Value', fontsize=25)
    plt.title('Box Plot - %s' % var, fontsize=40)
    plt.tight_layout()
    plt.savefig(r'D:\WUR\master\MLP\master thesis\data\数据汇总\descriptive_statistics\新\statistics_%s.png' % var)
    plt.show()

# ------------------------load 2 pc6 sites
# 注意，之所以是496，是因为删除了不对劲的那个小parcel以及dissolve by postcode!
PC6 = gpd.read_file(r'D:\WUR\master\MLP\master thesis\data\数据汇总\PC6_只有两个.geojson') # 496
site_1092 = gpd.read_file(r'D:\WUR\master\MLP\master thesis\data\数据汇总\PC6_1092.geojson') # 215
site_1094 = gpd.read_file(r'D:\WUR\master\MLP\master thesis\data\数据汇总\PC6_1094.geojson') # 281

# ------------------------totGVI
# Deeplabv3+
totGVI_DeepLab = gpd.read_file(r'D:\WUR\master\MLP\master thesis\data\数据汇总\totGVI\totGVI_DeepLab_PC6.geojson').\
    dissolve(by='Postcode',as_index=False)
totGVI_DeepLab = totGVI_DeepLab.drop(totGVI_DeepLab[totGVI_DeepLab['Postcode']=='1017TP'].index) # 496
totGVI_DeepLab_1092 = site_1092[['Postcode','geometry']].\
    merge(totGVI_DeepLab[['Postcode','totGVI']],how='left',left_on='Postcode',right_on='Postcode') # 215
totGVI_DeepLab_1094 = site_1094[['Postcode','geometry']].\
    merge(totGVI_DeepLab[['Postcode','totGVI']],how='left',left_on='Postcode',right_on='Postcode') # 281

totGVI_DeepLab = totGVI_DeepLab.rename(columns={'totGVI':'totGVI_DeepLab'})
totGVI_DeepLab_1092 = totGVI_DeepLab_1092.rename(columns={'totGVI':'totGVI_DeepLab'})
totGVI_DeepLab_1094 = totGVI_DeepLab_1094.rename(columns={'totGVI':'totGVI_DeepLab'})

create_box(totGVI_DeepLab_1092,totGVI_DeepLab_1094,totGVI_DeepLab,'totGVI_DeepLab')

# ------------------------top/bottom
# Deeplabv3+
VS_DeepLab = gpd.read_file(r'D:\WUR\master\MLP\master thesis\data\Vegetation_Structure\GSV_split\处理好的\VS_DeepLab_PC6.geojson').\
    dissolve(by='Postcode',as_index=False)
VS_DeepLab = VS_DeepLab.drop(VS_DeepLab[VS_DeepLab['Postcode']=='1017TP'].index) # 496
VS_DeepLab_1092 = site_1092[['Postcode','geometry']].\
    merge(VS_DeepLab[['Postcode','top','low']],how='left',left_on='Postcode',right_on='Postcode') # 215
VS_DeepLab_1094 = site_1094[['Postcode','geometry']].\
    merge(VS_DeepLab[['Postcode','top','low']],how='left',left_on='Postcode',right_on='Postcode') # 281

VS_DeepLab = VS_DeepLab.rename(columns={'top':'top_DeepLab','low':'low_DeepLab'})
VS_DeepLab_1092 = VS_DeepLab_1092.rename(columns={'top':'top_DeepLab','low':'low_DeepLab'})
VS_DeepLab_1094 = VS_DeepLab_1094.rename(columns={'top':'top_DeepLab','low':'low_DeepLab'})

create_box(VS_DeepLab_1092,VS_DeepLab_1094,VS_DeepLab,'top_DeepLab')
create_box(VS_DeepLab_1092,VS_DeepLab_1094,VS_DeepLab,'low_DeepLab')

# ------------------------VSD
PC6_VSD = gpd.read_file(r'D:\WUR\master\MLP\master thesis\data\数据汇总\VSD\PC6_species_diversity.geojson').\
    dissolve(by='Postcode',as_index=False)
PC6_VSD = PC6_VSD.drop(PC6_VSD[PC6_VSD['Postcode']=='1017TP'].index) # 496
VSD_1092 = site_1092[['Postcode','geometry']].\
    merge(PC6_VSD[['Postcode','SDI_indi']],how='left',left_on='Postcode',right_on='Postcode') # 215
VSD_1094 = site_1094[['Postcode','geometry']].\
    merge(PC6_VSD[['Postcode','SDI_indi']],how='left',left_on='Postcode',right_on='Postcode') # 281

PC6_VSD = PC6_VSD.rename(columns={'SDI_indi':'VSD'})
VSD_1092 = VSD_1092.rename(columns={'SDI_indi':'VSD'})
VSD_1094 = VSD_1094.rename(columns={'SDI_indi':'VSD'})

create_box(VSD_1092,VSD_1094,PC6_VSD,'VSD')