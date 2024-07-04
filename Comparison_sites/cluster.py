globals().clear()

import pandas as pd
import geopandas as gpd
import matplotlib.pyplot as plt
cluster_df = gpd.read_file(r'D:\WUR\master\MLP\master thesis\data\数据汇总\PC6_corr_df.geojson')[['totGVI', 'VSD', 'top', 'low',
       'vitality_level', 'noise_level','building_height', 'density', 'width_index','geometry','Postcode','avg_WI']]
cluster_df['noise_level'] = cluster_df['noise_level'].fillna(0)
cluster_df['vitality_level'] = cluster_df['vitality_level'].fillna(0)
print(cluster_df.columns)

width_mean = cluster_df['width_index'].mean()
cluster_df['width_index'] = cluster_df['width_index'].fillna(width_mean)

# 特征缩放
from sklearn import preprocessing
scaler = preprocessing.MinMaxScaler()
## 因为cluster analysis是在第二步，所以这里只对确认没有collinearity的变量作处理
columns_to_normalize = ['totGVI', 'VSD', 'top', 'low',
       'vitality_level', 'noise_level',
       'building_height', 'density', 'width_index']
scaler.fit(cluster_df[columns_to_normalize])
# Transform and replace the original columns with the scaled values
cluster_df[columns_to_normalize] = scaler.transform(cluster_df[columns_to_normalize])
x_norm = cluster_df.copy()[['totGVI', 'VSD', 'top', 'low',
       'vitality_level', 'noise_level',
       'building_height', 'density', 'width_index']]

# -----------------------k means
# Find the optimal number of clusters
from sklearn.cluster import KMeans
distortions = []
K = range(1,11)
for k in K:
 kmeanModel = KMeans(n_clusters=k)
 kmeanModel.fit(x_norm)
 distortions.append(kmeanModel.inertia_) # Plotting the distortions

plt.figure(figsize=(16,8))
plt.plot(K, distortions,marker='o')
plt.xlabel('k')
plt.ylabel('Distortion')
plt.title('The Elbow Method showing the optimal clusters')
plt.show()
plt.savefig(r'D:\WUR\master\MLP\master thesis\data\Comparison_sites\elbow_diagram.png')

# Put in the cluster num into our kmeans model
kmeans_model = KMeans(n_clusters=3, random_state=42)
kmeans_predict = kmeans_model.fit_predict(x_norm)
# Create a dataframe to store cluster assignments
cluster_df.loc[:,'Cluster'] = kmeans_predict
cluster_df.to_file(r'D:\WUR\master\MLP\master thesis\data\Comparison_sites\result_cluster.geojson',driver='GeoJSON')

cluster_summary = cluster_df.dissolve(by='Cluster',as_index=False,aggfunc={
                                                                           'totGVI':'mean', 'VSD':'mean', 'top':'mean', 'low':'mean',
                                                                           'vitality_level':'mean','noise_level':'mean',
    'building_height':'mean', 'density':'mean', 'width_index':'mean','avg_WI':'mean'})

##########
cluster_summary.to_csv(r'D:\WUR\master\MLP\master thesis\data\Comparison_sites\cluster_summary_original.csv', index=False)
a = pd.read_csv(r'D:\WUR\master\MLP\master thesis\data\Comparison_sites\cluster_summary_original.csv')

def transfer_word(df,doc_name):
    #pip install python-docx
    from docx import Document
    from docx.enum.table import WD_ALIGN_VERTICAL
    # Create a new Word document
    doc = Document()
    doc.add_paragraph('DataFrame Content:')
    table = doc.add_table(df.shape[0] + 1, df.shape[1])
    table.style = 'Table Grid'
    # Add headers to the first row
    for j in range(df.shape[1]):
        table.cell(0, j).text = df.columns[j]

    # Add the data from the DataFrame
    for i in range(df.shape[0]):
        for j in range(df.shape[1]):
            table.cell(i + 1, j).text = str(df.values[i, j])
    # Save the Word document
    doc.save(r"D:\WUR\master\MLP\master thesis\data\数据汇总\%s.docx"% doc_name)

transfer_word(a,'Cluster')

# 分别给每个pc4做一个饼图
cluster = gpd.read_file(r'D:\WUR\master\MLP\master thesis\data\Comparison_sites\result_cluster.geojson')
site_1092 = gpd.read_file(r'D:\WUR\master\MLP\master thesis\data\数据汇总\PC6_1092.geojson') # 215
site_1094 = gpd.read_file(r'D:\WUR\master\MLP\master thesis\data\数据汇总\PC6_1094.geojson') # 281
cluster_1092 = site_1092[['Postcode']].\
    merge(cluster[['Postcode','Cluster']],how='left',left_on='Postcode',right_on='Postcode')
cluster_1094 = site_1094[['Postcode']].\
    merge(cluster[['Postcode','Cluster']],how='left',left_on='Postcode',right_on='Postcode')

import pandas as pd
import matplotlib.pyplot as plt

cluster_counts = cluster_1092['Cluster'].value_counts()
# Custom colors for each cluster
custom_colors = {0: '#d3eaee', 1: '#1f78b4', 2: '#b2df8a'}
# Create a pie chart with a transparent background
fig, ax = plt.subplots(figsize=(8, 8))
ax.pie(cluster_counts, labels=cluster_counts.index, autopct='%1.1f%%', startangle=90,
       colors=[custom_colors.get(cluster, '#FFFFFF') for cluster in cluster_counts.index],  # Set custom colors
       wedgeprops=dict(width=0.4, edgecolor='w'),  # Control the width and edge color of wedges
       textprops=dict(color='black',fontsize=40))  # Set text color to white
# Set a transparent background
ax.set_facecolor('none')
# Set the title
plt.title('1092',fontsize=40)
plt.savefig(r'D:\WUR\master\MLP\master thesis\data\数据汇总\图\clustering\1092.png')

cluster_counts = cluster_1094['Cluster'].value_counts()
# Create a pie chart with a transparent background
fig, ax = plt.subplots(figsize=(8, 8))
ax.pie(cluster_counts, labels=cluster_counts.index, autopct='%1.1f%%', startangle=90,
       colors=[custom_colors.get(cluster, '#FFFFFF') for cluster in cluster_counts.index],  # Set custom colors
       wedgeprops=dict(width=0.4, edgecolor='w'),  # Control the width and edge color of wedges
       textprops=dict(color='black',fontsize=40))  # Set text color to white
# Set a transparent background
ax.set_facecolor('none')
# Set the title
plt.title('1094',fontsize=40)
plt.savefig(r'D:\WUR\master\MLP\master thesis\data\数据汇总\图\clustering\1094.png')