globals().clear()
import geopandas as gpd
import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns

def transfer_word(df,doc_name):
    #pip install python-docx
    from docx import Document
    from docx.enum.table import WD_ALIGN_VERTICAL
    # Create a new Word document
    doc = Document()
    doc.add_paragraph('DataFrame Content:')
    table = doc.add_table(df.shape[0] + 1, df.shape[1])
    table.style = 'Table Grid'
    # Add headers to the first row
    for j in range(df.shape[1]):
        table.cell(0, j).text = df.columns[j]

    # Add the data from the DataFrame
    for i in range(df.shape[0]):
        for j in range(df.shape[1]):
            table.cell(i + 1, j).text = str(df.values[i, j])
    # Save the Word document
    doc.save(r"D:\WUR\master\MLP\master thesis\data\数据汇总\comparison_analysis\%s.docx"% doc_name)

def create_violin(df1,df2,df3,var):
    # Calculate statistics for each site
    site1_stats = ['1092', df1[var].quantile(0.25), df1[var].quantile(0.5),df1[var].quantile(0.75),df1[var].mean()]
    site2_stats = ['1094', df2[var].quantile(0.25), df2[var].quantile(0.5), df2[var].quantile(0.75),df2[var].mean()]
    site3_stats = ['two sites', df3[var].quantile(0.25), df3[var].quantile(0.5), df3[var].quantile(0.75),df3[var].mean()]

    # Combine the statistics into a single DataFrame
    combined_stats = pd.DataFrame([site1_stats, site2_stats, site3_stats], columns=['Site', '25%', '50%', '75%','mean'])

    # Create a custom color palette for the violins
    #custom_palette = {'Site 1': 'blue', 'Site 2': 'green', 'Site 3': 'orange'}
    # Set the style using the custom palette
    #sns.set(style="whitegrid", palette=custom_palette)
    # Create a violin plot
    plt.figure(figsize=(10, 8))
    sns.violinplot(x='Site', y=var, data=pd.concat([df1.assign(Site='1092'), df2.assign(Site='1094'), df3.assign(Site='two sites')]))
    ax = sns.violinplot(x='Site', y=var,
                        data=pd.concat(
                            [df1.assign(Site='1092'), df2.assign(Site='1094'), df3.assign(Site='two sites')]))

    # Add lines and labels for the statistics
    for i, site in enumerate(combined_stats['Site']):
        plt.hlines(y=combined_stats.iloc[i, 1:4], xmin=i - 0.2, xmax=i + 0.2, color='red', linewidth=2)
        plt.text(i, ax.get_ylim()[0] - 0.03, f'Mean: {combined_stats.iloc[i, 4]:.2f}',
                 horizontalalignment='center', verticalalignment='bottom', color='red', fontsize=13)
        for j, stat in enumerate(['25%', '50%', '75%']):
            plt.text(i, combined_stats.iloc[i, j + 1], f'{stat}: {combined_stats.iloc[i, j + 1]:.2f}',
                     horizontalalignment='left', verticalalignment='bottom', color='red', fontsize=13)

    plt.xlabel('Site')
    plt.ylabel('Value')
    plt.title('Violin Plot with quantiles - %s'%var,fontsize=15)
    plt.tight_layout()
    #plt.savefig(r'D:\WUR\master\MLP\master thesis\data\数据汇总\descriptive_statistics\statistics_%s.png'%var)
    plt.savefig(r'D:\WUR\master\MLP\master thesis\data\数据汇总\descriptive_statistics\新\statistics_%s.png' % var)


# ------------------------load 2 pc6 sites
# 注意，之所以是496，是因为删除了不对劲的那个小parcel以及dissolve by postcode!
PC6 = gpd.read_file(r'D:\WUR\master\MLP\master thesis\data\数据汇总\PC6_只有两个.geojson') # 496
site_1092 = gpd.read_file(r'D:\WUR\master\MLP\master thesis\data\数据汇总\PC6_1092.geojson') # 215
site_1094 = gpd.read_file(r'D:\WUR\master\MLP\master thesis\data\数据汇总\PC6_1094.geojson') # 281

from statsmodels.stats.outliers_influence import variance_inflation_factor
import numpy as np
def identify_outliers(column):
    Q1 = np.percentile(column, 25)
    Q3 = np.percentile(column, 75)
    IQR = Q3 - Q1
    lower_bound = Q1 - 1.5 * IQR
    upper_bound = Q3 + 1.5 * IQR
    return (column < lower_bound) | (column > upper_bound)

# ---------------covariates
cor_df = gpd.read_file(r'D:\WUR\master\MLP\master thesis\data\数据汇总\PC6_corr_df.geojson').drop(columns='geometry')

cor_df['noise_level'] = cor_df['noise_level'].fillna(0)
cor_df['vitality_level'] = cor_df['vitality_level'].fillna(0)
print(cor_df.columns)

a = cor_df.columns[2:]
for col in a:
    outliers = identify_outliers(cor_df[col])
    cor_df.loc[outliers, col] = cor_df[col].mean()

cnesus_1092 = site_1092[['Postcode']].\
    merge(cor_df,how='left',left_on='Postcode',right_on='Postcode')
census_1094 = site_1094[['Postcode']].\
    merge(cor_df,how='left',left_on='Postcode',right_on='Postcode')

# --------------------------Mean/STD DEV calculateion
def calcu(df):
    df_1 = df.copy()
    df_1 = df_1.drop(columns=['Postcode'])
    for col in df_1.columns:
        print(col)
        # Mean for selected columns
        mean_name = 'mean' + '_' + col
        std_name = 'std_dev' + '_' + col
        df_1[mean_name] = df[col].mean()
        # Standard deviation for selected columns
        df_1[std_name] = df[col].std()
    df_1 = df_1.groupby(by=['mean_mean_NDVI','std_dev_mean_NDVI',
                            'mean_totGVI', 'std_dev_totGVI',
                            'mean_VSD','std_dev_VSD',
                            'mean_top', 'std_dev_top',
                            'mean_low', 'std_dev_low',
                            'mean_net_parking_pressure', 'std_dev_net_parking_pressure',
                            'mean_vitality_level', 'std_dev_vitality_level',
                            'mean_class_index','std_dev_class_index',
                            'mean_noise_level', 'std_dev_noise_level',
                            'mean_building_height', 'std_dev_building_height',
                            'mean_density', 'std_dev_density',
                            'mean_width_index', 'std_dev_width_index'],as_index=False).agg('first')
    df_1 = df_1[['mean_mean_NDVI','std_dev_mean_NDVI',
                            'mean_totGVI', 'std_dev_totGVI',
                            'mean_VSD','std_dev_VSD',
                            'mean_top', 'std_dev_top',
                            'mean_low', 'std_dev_low',
                            'mean_net_parking_pressure', 'std_dev_net_parking_pressure',
                 'mean_class_index','std_dev_class_index',
                            'mean_vitality_level', 'std_dev_vitality_level',
                            'mean_noise_level', 'std_dev_noise_level',
                            'mean_building_height', 'std_dev_building_height',
                            'mean_density', 'std_dev_density',
                            'mean_width_index', 'std_dev_width_index']]
    return df_1

cnesus_1092_calcu = calcu(cnesus_1092).T
cnesus_1092_calcu.columns = ['census_1092']
census_1094_calcu = calcu(census_1094).T
census_1094_calcu.columns = ['census_1094']

census_calcu = calcu(cor_df).T
census_calcu.columns = ['census']


description_stats = pd.concat([cnesus_1092_calcu,census_1094_calcu,
                               census_calcu],axis=1)
for col in description_stats.columns:
    description_stats[col] = description_stats[col].round(3)

description_stats['metrics'] = description_stats.index
description_stats.to_csv(r'D:\WUR\master\MLP\master thesis\data\数据汇总\comparison_analysis\description_stats.csv')
transfer_word(description_stats,'description_stats')

# ------------------------mean NDVI
# 493
NDVI = gpd.read_file(r'D:\WUR\master\MLP\master thesis\data\数据汇总\NDVI\meanNDVI_perPC6.geojson')
PC6_NDVI = PC6[['Postcode','geometry']].merge(NDVI[['Postcode','mean_NDVI']],how='left',left_on='Postcode',right_on='Postcode')
NDVI_1092 = site_1092[['Postcode','geometry']].\
    merge(NDVI[['Postcode','mean_NDVI']],how='left',left_on='Postcode',right_on='Postcode')
NDVI_1094 = site_1094[['Postcode','geometry']].\
    merge(NDVI[['Postcode','mean_NDVI']],how='left',left_on='Postcode',right_on='Postcode')

create_violin(NDVI_1092,NDVI_1094,PC6_NDVI,'mean_NDVI')

# ------------------------totGVI
# Deeplabv3+
totGVI_DeepLab = gpd.read_file(r'D:\WUR\master\MLP\master thesis\data\数据汇总\totGVI\totGVI_DeepLab_PC6.geojson').\
    dissolve(by='Postcode',as_index=False)
totGVI_DeepLab = totGVI_DeepLab.drop(totGVI_DeepLab[totGVI_DeepLab['Postcode']=='1017TP'].index) # 496
totGVI_DeepLab_1092 = site_1092[['Postcode','geometry']].\
    merge(totGVI_DeepLab[['Postcode','totGVI']],how='left',left_on='Postcode',right_on='Postcode') # 215
totGVI_DeepLab_1094 = site_1094[['Postcode','geometry']].\
    merge(totGVI_DeepLab[['Postcode','totGVI']],how='left',left_on='Postcode',right_on='Postcode') # 281

totGVI_DeepLab = totGVI_DeepLab.rename(columns={'totGVI':'totGVI_DeepLab'})
totGVI_DeepLab_1092 = totGVI_DeepLab_1092.rename(columns={'totGVI':'totGVI_DeepLab'})
totGVI_DeepLab_1094 = totGVI_DeepLab_1094.rename(columns={'totGVI':'totGVI_DeepLab'})

create_violin(totGVI_DeepLab_1092,totGVI_DeepLab_1094,totGVI_DeepLab,'totGVI_DeepLab')

# Pymeanshift
totGVI_pymean = gpd.read_file(r'D:\WUR\master\MLP\master thesis\data\数据汇总\totGVI\totGVI_Pymean_PC6.geojson').\
    dissolve(by='Postcode',as_index=False)
totGVI_pymean = totGVI_pymean.drop(totGVI_pymean[totGVI_pymean['Postcode']=='1017TP'].index) # 496
totGVI_pymean_1092 = site_1092[['Postcode','geometry']].\
    merge(totGVI_pymean[['Postcode','totGVI']],how='left',left_on='Postcode',right_on='Postcode') # 215
totGVI_pymean_1094 = site_1094[['Postcode','geometry']].\
    merge(totGVI_pymean[['Postcode','totGVI']],how='left',left_on='Postcode',right_on='Postcode') # 281

totGVI_pymean = totGVI_pymean.rename(columns={'totGVI':'totGVI_pymean'})
totGVI_pymean_1092 = totGVI_pymean_1092.rename(columns={'totGVI':'totGVI_pymean'})
totGVI_pymean_1094 = totGVI_pymean_1094.rename(columns={'totGVI':'totGVI_pymean'})

create_violin(totGVI_pymean_1092,totGVI_pymean_1094,totGVI_pymean,'totGVI_pymean')

# ------------------------VSD
PC6_VSD = gpd.read_file(r'D:\WUR\master\MLP\master thesis\data\数据汇总\VSD\PC6_species_diversity.geojson').\
    dissolve(by='Postcode',as_index=False)
PC6_VSD = PC6_VSD.drop(PC6_VSD[PC6_VSD['Postcode']=='1017TP'].index) # 496
VSD_1092 = site_1092[['Postcode','geometry']].\
    merge(PC6_VSD[['Postcode','SDI_indi']],how='left',left_on='Postcode',right_on='Postcode') # 215
VSD_1094 = site_1094[['Postcode','geometry']].\
    merge(PC6_VSD[['Postcode','SDI_indi']],how='left',left_on='Postcode',right_on='Postcode') # 281

PC6_VSD = PC6_VSD.rename(columns={'SDI_indi':'VSD'})
VSD_1092 = VSD_1092.rename(columns={'SDI_indi':'VSD'})
VSD_1094 = VSD_1094.rename(columns={'SDI_indi':'VSD'})

create_violin(VSD_1092,VSD_1094,PC6_VSD,'VSD')

# ------------------------top/middle/bottom
# Deeplabv3+
VS_DeepLab = gpd.read_file(r'D:\WUR\master\MLP\master thesis\data\数据汇总\VegetationStructure\NEW_VS_DeepLab_PC6.geojson').\
    dissolve(by='Postcode',as_index=False)
VS_DeepLab = VS_DeepLab.drop(VS_DeepLab[VS_DeepLab['Postcode']=='1017TP'].index) # 496
VS_DeepLab_1092 = site_1092[['Postcode','geometry']].\
    merge(VS_DeepLab[['Postcode','top','low']],how='left',left_on='Postcode',right_on='Postcode') # 215
VS_DeepLab_1094 = site_1094[['Postcode','geometry']].\
    merge(VS_DeepLab[['Postcode','top','low']],how='left',left_on='Postcode',right_on='Postcode') # 281

VS_DeepLab = VS_DeepLab.rename(columns={'top':'top_DeepLab','low':'low_DeepLab'})
VS_DeepLab_1092 = VS_DeepLab_1092.rename(columns={'top':'top_DeepLab','low':'low_DeepLab'})
VS_DeepLab_1094 = VS_DeepLab_1094.rename(columns={'top':'top_DeepLab','low':'low_DeepLab'})

def create_violin_low(df1,df2,df3,var):
    # Calculate statistics for each site
    site1_stats = ['1092', df1[var].quantile(0.25), df1[var].quantile(0.5),df1[var].quantile(0.75),df1[var].mean()]
    site2_stats = ['1094', df2[var].quantile(0.25), df2[var].quantile(0.5), df2[var].quantile(0.75),df2[var].mean()]
    site3_stats = ['two sites', df3[var].quantile(0.25), df3[var].quantile(0.5), df3[var].quantile(0.75),df3[var].mean()]

    # Combine the statistics into a single DataFrame
    combined_stats = pd.DataFrame([site1_stats, site2_stats, site3_stats], columns=['Site', '25%', '50%', '75%','mean'])

    plt.figure(figsize=(10, 8))
    sns.violinplot(x='Site', y=var, data=pd.concat([df1.assign(Site='1092'), df2.assign(Site='1094'), df3.assign(Site='two sites')]))
    ax = sns.violinplot(x='Site', y=var,
                        data=pd.concat(
                            [df1.assign(Site='1092'), df2.assign(Site='1094'), df3.assign(Site='two sites')]))

    # Add lines and labels for the statistics
    for i, site in enumerate(combined_stats['Site']):
        plt.hlines(y=combined_stats.iloc[i, 1:4], xmin=i - 0.2, xmax=i + 0.2, color='red', linewidth=2)
        plt.text(i, ax.get_ylim()[0] - 0.03, f'Mean: {combined_stats.iloc[i, 4]:.2f}',
                 horizontalalignment='center', verticalalignment='bottom', color='red', fontsize=13)
        for j, stat in enumerate(['25%', '50%', '75%']):
            if stat == '25%':
                plt.text(i, ax.get_ylim()[0] + 2.5, f'{stat}: {combined_stats.iloc[i, j + 1]:.2f}',
                         horizontalalignment='left', verticalalignment='bottom', color='red', fontsize=13)
            else:
                plt.text(i, combined_stats.iloc[i, j + 1], f'{stat}: {combined_stats.iloc[i, j + 1]:.2f}',
                     horizontalalignment='left', verticalalignment='bottom', color='red', fontsize=13)

    plt.xlabel('Site')
    plt.ylabel('Value')
    plt.title('Violin Plot with quantiles - %s'%var,fontsize=15)
    plt.tight_layout()
    plt.savefig(r'D:\WUR\master\MLP\master thesis\data\数据汇总\descriptive_statistics\statistics_%s.png'%var)

create_violin(VS_DeepLab_1092,VS_DeepLab_1094,VS_DeepLab,'top_DeepLab')
create_violin(VS_DeepLab_1092,VS_DeepLab_1094,VS_DeepLab,'low_DeepLab')

# Pymeanshift
VS_pymean = gpd.read_file(r'D:\WUR\master\MLP\master thesis\data\数据汇总\VegetationStructure\VS_Pymean_PC6.geojson').\
    dissolve(by='Postcode',as_index=False)
VS_pymean = VS_pymean.drop(VS_pymean[VS_pymean['Postcode']=='1017TP'].index) # 496
VS_pymean_1092 = site_1092[['Postcode','geometry']].\
    merge(VS_pymean[['Postcode','top','middle','low']],how='left',left_on='Postcode',right_on='Postcode') # 215
VS_pymean_1094 = site_1094[['Postcode','geometry']].\
    merge(VS_pymean[['Postcode','top','middle','low']],how='left',left_on='Postcode',right_on='Postcode') # 281

VS_pymean = VS_pymean.rename(columns={'top':'top_pymean','middle':'middle_pymean','low':'low_pymean'})
VS_pymean_1092 = VS_pymean_1092.rename(columns={'top':'top_pymean','middle':'middle_pymean','low':'low_pymean'})
VS_pymean_1094 = VS_pymean_1094.rename(columns={'top':'top_pymean','middle':'middle_pymean','low':'low_pymean'})

create_violin(VS_pymean_1092,VS_pymean_1094,VS_pymean,'top_pymean')
create_violin(VS_pymean_1092,VS_pymean_1094,VS_pymean,'middle_pymean')
create_violin_low(VS_pymean_1092,VS_pymean_1094,VS_pymean,'low_pymean')

# ------------------------enclosure_level
# Sample ordinal data (replace with your data)
PC6_EL = gpd.read_file(r'D:\WUR\master\MLP\master thesis\data\数据汇总\EnclosureLevel\PC6_enclosrue_level_cluster.geojson')[['Postcode','Cluster']]
# 496
# 还是不要赋值了，感觉有点不对
for idx, row in PC6_EL.iterrows():
    if row['Cluster'] == 0:
        PC6_EL.at[idx,'Cluster'] = 1
    elif row['Cluster'] == 1:
        PC6_EL.at[idx, 'Cluster'] = 2
    elif row['Cluster'] == 2:
        PC6_EL.at[idx, 'Cluster'] = 3
    elif row['Cluster'] == 3:
        PC6_EL.at[idx, 'Cluster'] = 4

EL_1092 = site_1092[['Postcode','geometry']].\
    merge(PC6_EL[['Postcode','Cluster']],how='left',left_on='Postcode',right_on='Postcode') # 215
EL_1094 = site_1094[['Postcode','geometry']].\
    merge(PC6_EL[['Postcode','Cluster']],how='left',left_on='Postcode',right_on='Postcode') # 281

def frequency_calcu(df):
    # Create a frequency table
    frequency_table = df['Cluster'].value_counts().reset_index()
    # Rename columns for clarity
    frequency_table.columns = ['Cluster', 'Frequency']
    # Sort the table by category if needed
    frequency_table = frequency_table.sort_values(by='Cluster')
    # Display the frequency table
    #print(frequency_table)
    return frequency_table

EL_1092_frequency = frequency_calcu(EL_1092)
EL_1094_frequency = frequency_calcu(EL_1094)
PC6_EL_frequency = frequency_calcu(PC6_EL)
EL_table = pd.DataFrame(columns = ['Cluster','1092','1094','all'])
EL_table['Cluster'] = EL_1092_frequency['Cluster']

EL_table['1092'] = EL_1092_frequency['Frequency']
EL_table['1092_perc'] = EL_table['1092']/sum(EL_table['1092'].values)
EL_table['1094'] = EL_1094_frequency['Frequency']
EL_table['1094'] = EL_table['1094'].fillna(0)
EL_table['1094_perc'] = EL_table['1094']/sum(EL_table['1094'].values)
EL_table['all'] = PC6_EL_frequency['Frequency']
EL_table['all_perc'] = EL_table['all']/sum(EL_table['all'].values)

for col in EL_table.columns:
    EL_table[col] = EL_table[col].round(3)

EL_table['Cluster'] = EL_table['Cluster'].astype('int')
EL_table.to_csv(r'D:\WUR\master\MLP\master thesis\data\数据汇总\descriptive_statistics\EL_table.csv')
transfer_word(EL_table,'frequency_table')

# 这一段没什么用，不要用这种表现形式，就pie chart最好
"""
# plot three figures in one row: 1092, 1094, all
fig, axes = plt.subplots(1, 3, figsize=(15, 5))
# Reset to the default color palette
sns.set_palette("Paired")
# Plot the second bar plot in the second subplot
sns.barplot(y=EL_table['1092_perc'], x=EL_table['EL'], ax=axes[0])
axes[0].set_title('1092')
axes[0].set_ylabel('Values')

sns.barplot(y=EL_table['1094_perc'], x=EL_table['EL'], ax=axes[1])
axes[1].set_title('1094')
axes[1].set_ylabel('Values')

sns.barplot(y=EL_table['all_perc'], x=EL_table['EL'], ax=axes[2])
axes[2].set_title('two sites')
axes[2].set_ylabel('Values')

# Add a common legend for all subplots
handles, labels = axes[0].get_legend_handles_labels()
fig.legend(handles, labels, loc='upper right')
fig.suptitle('Enclosure level for sites')
plt.tight_layout()
# Show the combined plot
plt.show()
plt.savefig(r'D:\WUR\master\MLP\master thesis\data\数据汇总\comparison_analysis\enclosure_level_comparison.png')
"""

# -------------------------------pie charts for enclosure level
# plot three figures in one row: 1092, 1094, all
fig, axes = plt.subplots(1, 3, figsize=(15, 5))
# First Pie Chart
plt.subplot(131)  # 1 row, 3 columns, position 1
plt.pie(EL_table['1092_perc'], labels=EL_table['Cluster'], autopct='%1.1f%%', startangle=90)
plt.title('1092_enclosure_level_proportion')
plt.subplot(132)  # 1 row, 3 columns, position 1
plt.pie(EL_table['1094_perc'], labels=EL_table['Cluster'], autopct='%1.1f%%', startangle=90)
plt.title('1094_enclosure_level_proportion')
plt.subplot(133)  # 1 row, 3 columns, position 1
plt.pie(EL_table['all_perc'], labels=EL_table['Cluster'], autopct='%1.1f%%', startangle=90)
plt.title('all sites')

# Add a common legend for all subplots
handles, labels = axes[0].get_legend_handles_labels()
fig.legend(handles, labels, loc='upper right')
fig.suptitle('Cluster types for sites')
plt.tight_layout()
# Show the combined plot
plt.show()
plt.savefig(r'D:\WUR\master\MLP\master thesis\data\数据汇总\comparison_analysis\enclosure_level_comparison.png')

# ------------------------Walkability Index (PC6)
PC6_walkabilityindex = gpd.read_file(r'D:\WUR\master\MLP\master thesis\data\数据汇总\PC6_WalkabilityIndex.geojson')[
    ['Postcode', 'geometry', 'WI7c_1719_150m_znorm']]
PC6_walkabilityindex = PC6_walkabilityindex.rename(columns={'WI7c_1719_150m_znorm': 'WalkabilityIndex'})
PC6_walkabilityindex = PC6_walkabilityindex.drop(
    PC6_walkabilityindex[PC6_walkabilityindex['Postcode'] == '1017TP'].index)
# 496
walkabilityindex_1092 = site_1092[['Postcode','geometry']].\
    merge(PC6_walkabilityindex[['Postcode','WalkabilityIndex']],how='left',left_on='Postcode',right_on='Postcode') # 215
walkabilityindex_1094 = site_1094[['Postcode','geometry']].\
    merge(PC6_walkabilityindex[['Postcode','WalkabilityIndex']],how='left',left_on='Postcode',right_on='Postcode') # 281

create_violin(walkabilityindex_1092,walkabilityindex_1094,PC6_walkabilityindex,'WalkabilityIndex')

# -----------------------组合1092/1094/tot的数据
## -----------1092
# Deeplabv3+
dfs = [walkabilityindex_1092, NDVI_1092.drop(columns='geometry'), totGVI_DeepLab_1092.drop(columns='geometry'),
       VSD_1092.drop(columns='geometry'), VS_DeepLab_1092.drop(columns='geometry')]
from functools import reduce
table_deeplab_1092 = reduce(lambda left, right: pd.merge(left, right, on=['Postcode'],
                                             how='left'), dfs)
# Pymeanshift
dfs = [walkabilityindex_1092, NDVI_1092.drop(columns='geometry'), totGVI_pymean_1092.drop(columns='geometry'),
       VSD_1092.drop(columns='geometry'), VS_pymean_1092.drop(columns='geometry')]
from functools import reduce
table_pymean_1092 = reduce(lambda left, right: pd.merge(left, right, on=['Postcode'],
                                             how='left'), dfs)

## -----------1094
dfs = [walkabilityindex_1094, NDVI_1094.drop(columns='geometry'), totGVI_DeepLab_1094.drop(columns='geometry'),
       VSD_1094.drop(columns='geometry'), VS_DeepLab_1094.drop(columns='geometry')]
from functools import reduce
table_deeplab_1094 = reduce(lambda left, right: pd.merge(left, right, on=['Postcode'],
                                             how='left'), dfs)
# Pymeanshift
dfs = [walkabilityindex_1094, NDVI_1094.drop(columns='geometry'), totGVI_pymean_1094.drop(columns='geometry'),
       VSD_1094.drop(columns='geometry'), VS_pymean_1094.drop(columns='geometry'), EL_1094.drop(columns='geometry')]
from functools import reduce
table_pymean_1094 = reduce(lambda left, right: pd.merge(left, right, on=['Postcode'],
                                             how='left'), dfs)

## -----------tot
dfs = [PC6_walkabilityindex, PC6_NDVI.drop(columns='geometry'),
       totGVI_DeepLab.drop(columns='geometry'), PC6_VSD.drop(columns='geometry'),
       VS_DeepLab.drop(columns='geometry')]
table_deeplab_PC6 = reduce(lambda left, right: pd.merge(left, right, on=['Postcode'],
                                             how='left'), dfs)
table_deeplab_PC6_scaled = gpd.read_file(r'D:\WUR\master\MLP\master thesis\data\数据汇总\PC6_corr_df.geojson')

dfs = [PC6_walkabilityindex, PC6_NDVI.drop(columns='geometry'),
       totGVI_pymean.drop(columns='geometry'), PC6_VSD.drop(columns='geometry'),
       VS_pymean.drop(columns='geometry'), PC6_EL]
table_pymean_PC6 = reduce(lambda left, right: pd.merge(left, right, on=['Postcode'],
                                             how='left'), dfs).\
    drop(columns=['top_right', 'top_middle', 'top_left', 'middle_right',
       'middle_middle', 'middle_left', 'low_right', 'low_middle', 'low_left'])

# --------------------------Mean/STD DEV calculateion
def calcu(df):
    df_1 = df.copy()
    df_1 = df_1.drop(columns=['Postcode','WalkabilityIndex','geometry','EL']) # 1092/1094只有EL没有Cluster
    for col in df_1.columns:
        print(col)
        # Mean for selected columns
        mean_name = 'mean' + '_' + col
        std_name = 'std_dev' + '_' + col
        df_1[mean_name] = df[col].mean()
        # Standard deviation for selected columns
        df_1[std_name] = df[col].std()
    df_1 = df_1.groupby(by=['mean_mean_NDVI',
       'std_dev_mean_NDVI', 'mean_totGVI', 'std_dev_totGVI', 'mean_SDI_indi',
       'std_dev_SDI_indi', 'mean_top', 'std_dev_top', 'mean_middle',
       'std_dev_middle', 'mean_low', 'std_dev_low'],as_index=False).agg('first')
    df_1 = df_1[['mean_mean_NDVI',
       'std_dev_mean_NDVI', 'mean_totGVI', 'std_dev_totGVI', 'mean_SDI_indi',
       'std_dev_SDI_indi', 'mean_top', 'std_dev_top', 'mean_middle',
       'std_dev_middle', 'mean_low', 'std_dev_low']]
    return df_1

table_pymean_PC6_calcu = calcu(table_pymean_PC6).T
table_pymean_PC6_calcu.columns = ['pymean_PC6']
table_deeplab_PC6_calcu = calcu(table_deeplab_PC6).T
table_deeplab_PC6_calcu.columns = ['deeplab_PC6']

table_pymean_1092_calcu = calcu(table_pymean_1092).T
table_pymean_1092_calcu.columns = ['pymean_1092']
table_pymean_1094_calcu = calcu(table_pymean_1094).T
table_pymean_1094_calcu.columns = ['pymean_1094']

table_DeepLab_1092_calcu = calcu(table_deeplab_1092).T
table_DeepLab_1092_calcu.columns = ['pymean_1092']
table_DeepLab_1094_calcu = calcu(table_deeplab_1094).T
table_DeepLab_1094_calcu.columns = ['pymean_1094']

description_stats = pd.concat([table_pymean_PC6_calcu,table_deeplab_PC6_calcu,
                               table_pymean_1092_calcu,table_pymean_1094_calcu,
                               table_DeepLab_1092_calcu,table_DeepLab_1094_calcu],axis=1)
for col in description_stats.columns:
    description_stats[col] = description_stats[col].round(3)

description_stats['metrics'] = description_stats.index
description_stats.to_csv(r'D:\WUR\master\MLP\master thesis\data\数据汇总\comparison_analysis\description_stats.csv')
transfer_word(description_stats,'description_stats')

# walkability index statistics
WI = pd.DataFrame(columns=['mean_WalkabilityIndex_1092','std_WalkabilityIndex_1092',
                           'mean_WalkabilityIndex_1094','std_WalkabilityIndex_1094',
                           'mean_WalkabilityIndex_tot','std_WalkabilityIndex_tot'])
WI.loc[0,'mean_WalkabilityIndex_1092'] = walkabilityindex_1092['WalkabilityIndex'].mean()
WI.loc[0,'std_WalkabilityIndex_1092'] = walkabilityindex_1092['WalkabilityIndex'].std()
WI.loc[0,'mean_WalkabilityIndex_1094'] = walkabilityindex_1094['WalkabilityIndex'].mean()
WI.loc[0,'std_WalkabilityIndex_1094'] = walkabilityindex_1094['WalkabilityIndex'].std()
WI.loc[0,'mean_WalkabilityIndex_tot'] = PC6_walkabilityindex['WalkabilityIndex'].mean()
WI.loc[0,'std_WalkabilityIndex_tot'] = PC6_walkabilityindex['WalkabilityIndex'].std()

WI = WI.round(3)

WI.to_csv(r'D:\WUR\master\MLP\master thesis\data\数据汇总\comparison_analysis\WalkabilityIndex_statcs.csv')
transfer_word(WI,'WalkabilityIndex_statcs')

# --------------------------semantic segmentation comparison
import matplotlib.pyplot as plt
import numpy as np
table_pymean = table_pymean_PC6[['totGVI','top','middle','low']]
table_deeplab = table_deeplab_PC6[['totGVI','top','middle','low']]

plt.figure(figsize=(20,8))
plt.plot(table_pymean[['totGVI']], label='pymean', marker='o',linewidth=0.5)
plt.plot(table_deeplab[['totGVI']], label='deeplab', marker='s',linewidth=0.5)
plt.xlabel('PC6s')
plt.ylabel('totGVI measures')
plt.legend()
plt.title('Comparison of Two Methods')
plt.show()
plt.savefig(r'D:\WUR\master\MLP\master thesis\data\数据汇总\comparison_analysis\totGVI_comparison.png')

# ---------------------violin plots for comparison
stats = pd.read_csv(r'D:\WUR\master\MLP\master thesis\data\数据汇总\comparison_analysis\Discussion_porprotion.csv').iloc[:,0:5]
# Create a bar plot to compare means
import matplotlib.pyplot as plt
print(stats.columns)
variables = ['1092_mean', '1094_mean', '1092_std','1094_std']

# Define colors for each column
colors = ['red', 'green', 'blue', 'orange']

for i in range(len(stats)):
    print(i)
    plt.figure(figsize=(8, 4))  # Set the figure size for each chart
    categories = ['1092_mean', '1094_mean', '1092_std','1094_std']
    values = stats.iloc[i, 1:]
    title = stats.iloc[i, 0]
    # Use the 'colors' list to set different colors for each bar
    plt.bar(categories, values, color=colors)
    plt.xlabel('mean and standard deviations',fontsize=13)
    plt.ylabel('Values',fontsize=13)
    plt.title(f'Bar Chart for %s'%title,fontsize=16)
    plt.show()
    plt.savefig(r'D:\WUR\master\MLP\master thesis\data\数据汇总\comparison_analysis\%s.png'%title)


cor_Df_scaled = pd.read_csv(r'D:\WUR\master\MLP\master thesis\data\数据汇总\PC6_corr_df.csv')
reference_scaled_mean = pd.DataFrame(columns=cor_Df_scaled.columns)
for col in reference_scaled_mean.columns:
    reference_scaled_mean.loc[0,col] = cor_Df_scaled[col].mean().round(3)
reference_scaled_mean.to_csv(r'D:\WUR\master\MLP\master thesis\data\数据汇总\scaled数据mean值对照表.csv')