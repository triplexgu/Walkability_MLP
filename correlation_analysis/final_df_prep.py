globals().clear()
import pandas as pd

# ------------导入GAM summary
import re
# Specify the path to the GAM model summary text file
file_path = "D:\WUR\master\MLP\master thesis\data\数据汇总\Corr_analysis\GAM_summary.txt"
with open(file_path, "r", encoding="utf-8") as file:  # Specify the correct encoding (e.g., utf-8)
    lines = file.readlines()
# Extract the lines containing the summary information
summary_lines = [line.strip() for line in lines if re.match(r"^[A-Za-z]", line)]
# Extract the headers (variable names)
headers = ["Estimate", "edf", "Ref.df", "F",'p-value','significance']
# Extract the data rows
data = [line.split() for line in summary_lines[7:16]]

# Create a DataFrame
gam_summary_df = pd.DataFrame(data, columns=headers)
# Convert numeric columns to appropriate data types (e.g., float)
#numeric_columns = ["Estimate", "Std. Error", "t value", "Pr(>|t|)"]
#gam_summary_df[numeric_columns] = gam_summary_df[numeric_columns].astype(float)
# Print the DataFrame
#print(gam_summary_df)
gam_summary_df.to_csv(r'D:\WUR\master\MLP\master thesis\data\数据汇总\Corr_analysis\gam_summary.csv')

# ------------整理GAM summary
gam_summary_df = pd.read_csv(r'D:\WUR\master\MLP\master thesis\data\数据汇总\Corr_analysis\gam_summary.csv')

# ------------导入OLS summary
ols_summary_df = pd.read_csv(r'D:\WUR\master\MLP\master thesis\data\数据汇总\Corr_analysis\ols_summary_table.csv')
for col in ols_summary_df.columns[1:]:
    ols_summary_df[col] = ols_summary_df[col].round(3)

# ------------导入RF summary
rf_summary_df = pd.read_csv(r'D:\WUR\master\MLP\master thesis\data\数据汇总\Corr_analysis\rf_summary_table.csv')
rf_summary_df['Importance'] = rf_summary_df['Importance'].round(3)
# visalize feature importance
# Create the bar chart
import seaborn as sns
import matplotlib.pyplot as plt
plt.figure(figsize=(10, 6))
sns.barplot(x='Importance', y='Feature', data=rf_summary_df, orient='h', palette='viridis')
# Customize the chart (labels, title, etc.)
plt.xlabel('Importance')
plt.ylabel('Feature')
plt.title('Feature Importance Bar Chart')
# Show the chart
sns.plt.show()
plt.savefig(r'D:\WUR\master\MLP\master thesis\data\数据汇总\Corr_analysis\feature_importance.png', bbox_inches='tight')

# 把csv转化成word
def transfer_word(df,doc_name):
    #pip install python-docx
    from docx import Document
    from docx.enum.table import WD_ALIGN_VERTICAL
    # Create a new Word document
    doc = Document()
    doc.add_paragraph('DataFrame Content:')
    table = doc.add_table(df.shape[0] + 1, df.shape[1])
    table.style = 'Table Grid'
    # Add headers to the first row
    for j in range(df.shape[1]):
        table.cell(0, j).text = df.columns[j]

    # Add the data from the DataFrame
    for i in range(df.shape[0]):
        for j in range(df.shape[1]):
            table.cell(i + 1, j).text = str(df.values[i, j])
    # Save the Word document
    doc.save(r"D:\WUR\master\MLP\master thesis\data\数据汇总\Corr_analysis\%s.docx"% doc_name)

transfer_word(rf_summary_df,'rf_summary')
transfer_word(gam_summary_df,'gam_summary')
transfer_word(ols_summary_df,'ols_summary')